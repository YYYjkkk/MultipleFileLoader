from typing import List
import tqdm
from langchain_community.document_loaders.unstructured import UnstructuredFileLoader


class RapidOCRDocLoader(UnstructuredFileLoader):
    def _get_elements(self) -> List:
        def doc2text(filepath):
            from io import BytesIO
            import numpy as np
            from docx import Document, ImagePart
            from docx.oxml.table import CT_Tbl
            from docx.oxml.text.paragraph import CT_P
            from docx.table import Table, _Cell
            from docx.text.paragraph import Paragraph
            from PIL import Image
            from rapidocr_onnxruntime import RapidOCR

            ocr = RapidOCR()
            doc = Document(filepath)
            resp = ""

            def iter_block_items(parent):
                from docx.document import Document

                if isinstance(parent, Document):
                    parent_elm = parent.element.body
                elif isinstance(parent, _Cell):
                    parent_elm = parent._tc
                else:
                    raise ValueError("RapidOCRDocLoader parse fail")

                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P):
                        yield Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        yield Table(child, parent)

            b_unit = tqdm.tqdm(
                total=len(doc.paragraphs) + len(doc.tables),
                desc="RapidOCRDocLoader block index: 0",
            )
            for i, block in enumerate(iter_block_items(doc)):
                b_unit.set_description("RapidOCRDocLoader block index: {}".format(i))
                b_unit.refresh()
                if isinstance(block, Paragraph):
                    resp += block.text.strip() + "\n"
                    images = block._element.xpath(".//pic:pic")  # Get all images
                    for image in images:
                        for img_id in image.xpath(".//a:blip/@r:embed"):  # Get image id
                            part = doc.part.related_parts[img_id]  # Get corresponding image by id
                            if isinstance(part, ImagePart):
                                image = Image.open(BytesIO(part._blob))
                                result, _ = ocr(np.array(image))
                                if result:
                                    ocr_result = [line[1] for line in result]
                                    resp += "\n".join(ocr_result)
                elif isinstance(block, Table):
                    for row in block.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                resp += paragraph.text.strip() + "\n"
                b_unit.update(1)
            return resp

        text = doc2text(self.file_path)
        from unstructured.partition.text import partition_text
        return partition_text(text=text, **self.unstructured_kwargs)

    def save_to_txt(self, output_path: str):
        docs = self.load()
        with open(output_path, 'w', encoding='utf-8') as f:
            for doc in docs:
                f.write(doc.page_content + "\n")  # Write each document to the text file


if __name__ == "__main__":
    loader = RapidOCRDocLoader(file_path="document_loaders/input/testForDocx.docx")
    
    # Specify the output path for the .txt file
    output_txt_path = "document_loaders/output/docx2txt_output.txt"
    
    loader.save_to_txt(output_txt_path)
    print(f"Converted text saved to {output_txt_path}")
